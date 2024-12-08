import os
import io
import docx
import streamlit as st
import re
import time
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from datetime import datetime, timedelta
from typing import List, Dict, Union, IO, Optional
import markdown
from bs4 import BeautifulSoup
from streamlit_extras.switch_page_button import switch_page
from streamlit_cookies_manager import CookieManager
import pandas as pd
import tiktoken
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from graphviz import Source
from pyvis.network import Network
import tempfile

### CONSTANTS ###

STREAMLIT_KEY_NAME = "TransactionMonitoring_credential_ACN"

### STREAMLIT UTILITIES ###

def show_logo(image_path: str = 'code/images/logo.png'):
    st.logo(image_path)
    st.html("<style> [alt=Logo] {height: 4rem;} </style>")

def show_selected_case():
    if 'selected_case' in st.session_state:
        selected_case = st.session_state['selected_case']
    else:
        selected_case = None
    st.sidebar.markdown(f"Selected case: **{selected_case}**")
    if st.sidebar.button("Log Out"):
        st.session_state["delete_password"] = True
        switch_page("transaction monitoring")

def alert_and_redirect():
    st.error("There is no case selected. Please go to the **Case selector** page to select a case.")
    time.sleep(2)
    st.toast("Redirecting to case selector...", icon="🔄")
    time.sleep(1)
    switch_page("case selector")

### FILE AND DATA OPERATIONS ###

def get_docx_text(file: Union[str, bytes, bytearray, memoryview]) -> str:
    """
    Extract text from a docx file, including paragraphs and tables in order.
    """
    file_obj: Union[str, IO[bytes]]
    
    if isinstance(file, (bytes, bytearray, memoryview)):
        # Convert bytearray and memoryview to bytes
        if isinstance(file, (bytearray, memoryview)):
            file = bytes(file)
        file_obj = io.BytesIO(file)
    else:
        file_obj = file

    doc = docx.Document(file_obj)
    fullText = []

    for element in doc.element.body:
        if isinstance(element, CT_P):  # Check if the element is a paragraph
            para = docx.text.paragraph.Paragraph(element, doc)
            fullText.append(para.text)
        elif isinstance(element, CT_Tbl):  # Check if the element is a table
            table = docx.table.Table(element, doc)
            md_table = table_to_markdown(table)
            fullText.append(md_table)
            fullText.append('\n')  # Separate tables with a newline

    return '\n'.join(fullText)

def table_to_markdown(table: docx.table.Table) -> str:
    """
    Convert a docx table to a Markdown formatted string.
    """
    md_table = []
    # Extract headers
    headers = []
    for cell in table.rows[0].cells:
        headers.append(cell.text.strip())
    md_table.append('| ' + ' | '.join(headers) + ' |')
    md_table.append('|' + ' --- |' * len(headers))

    # Extract rows
    for row in table.rows[1:]:
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text.strip())
        md_table.append('| ' + ' | '.join(row_text) + ' |')
    
    return '\n'.join(md_table)

def get_folders(path: str) -> List[str]:
    return [name for name in os.listdir(path) if os.path.isdir(os.path.join(path, name))]

def get_files_in_folder(folder_path: str) -> List[str]:
    return [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

def get_files_in_folder_recursive(folder_path: str) -> List[str]:
    files = []
    for root, _, filenames in os.walk(folder_path):
        for filename in filenames:
            files.append(os.path.join(root, filename))
    return files

def list_cases(path: str, filter_by_word: str, full_path: bool = True, remove_items: List[str] = ['C0 - Información Común']) -> List[str]:
    """
    List all cases (folders only) in the given path, filter the list by a contained word and remove specific items. 
    """
    folders = [name for name in os.listdir(path) if os.path.isdir(os.path.join(path, name))]
    if full_path:
        return [os.path.join(path, folder) for folder in folders]
    if filter_by_word and filter_by_word != "ALL":
        folders = [folder for folder in folders if filter_by_word in folder]
    if remove_items:
        folders = [folder for folder in folders if folder not in remove_items]
    return folders

@st.cache_data
def read_json_from_blob(_blob_client, case: str, folder: str, json_file: str) -> Dict:
    return json.loads(_blob_client.get_file(os.path.join(folder, case, json_file)))

def load_json(
    case: str,
    keywords: List[str],
    files: List[str],
    folder: str,
    _blob_client,
    height=300
) -> Optional[List[Dict]]:
    """ Finds and displays a JSON file based on a keyword in the filename. """
    with st.container(border=True):
        json_files = [f for f in files if any(k in f for k in keywords)]
        if json_files:
            selected_json_files = st.multiselect(f"Selected {keywords[0]} JSON file", options=json_files, default=json_files)
            selected_json_data = []
            with st.container(border=False, height=height):
                for json_file in selected_json_files:
                    # json_data = read_json(case, json_file)
                    json_data = read_json_from_blob(
                        _blob_client=_blob_client,
                        case=case,
                        folder=folder,
                        json_file=json_file
                    )
                    selected_json_data.append(json_data)
                    st.json(json_data)
                    st.divider()
            return selected_json_data

def get_response_filepath(folder_path: str = 'narratives/'):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    file_name = st.session_state['selected_case'] + "_narrative.txt"
    file_path = os.path.join(folder_path, file_name)
    return file_path


def write_to_file(response: str, file_path: str):
    with open(file_path, 'a') as f:
        f.write(response)

# @st.cache_data
def read_excel_with_dynamic_headers(filepath, search_limit_rows=10, sheet=None):
    """
    Reads an Excel file, identifying the header row based on content and maximum column criteria.
    
    :param filepath: Path to the Excel file.
    :param search_limit_rows: The maximum number of rows to search for the header.
    :return: A pandas DataFrame with the correct headers.
    """
    # Step 1: Open the Excel file without headers
    if sheet:
        df = pd.read_excel(filepath, header=None, sheet_name=sheet)
    else:
        df = pd.read_excel(filepath, header=None)
    # Step 2: Identify the maximum number of content columns
    max_content_columns = df.apply(lambda x: x.count(), axis=1).max()
    
    # Step 3: Iterate through rows to find the header
    header_row_idx = None
    for i, row in df.iterrows():
        if row.isna().all():
            continue  # Skip completely NaN rows
        elif row.count() == max_content_columns:
            # Found a row that matches the maximum number of content columns
            header_row_idx = i
            break
        elif i >= search_limit_rows:
            break
    
    if header_row_idx is not None:
        # Step 4: Re-read the Excel file with the identified header row
        if sheet:
            df = pd.read_excel(filepath, header=header_row_idx, sheet_name=sheet)
        else:
            df = pd.read_excel(filepath, header=header_row_idx)
    
    return df

def read_additional_documentation(narrative_path: str, case: str) -> Dict[str, str]:
    files = get_files_in_folder_recursive(os.path.join(narrative_path, case))
    files_content = {}
    for file in files:
        if os.path.splitext(file)[1] == ".docx":
            docx_file = get_docx_text(file)
            files_content[os.path.basename(file)] = docx_file
    return files_content

def md_to_text(md: str) -> str:
    html = markdown.markdown(md)
    soup = BeautifulSoup(html, features='html.parser')
    return soup.get_text()

### VISUALIZATION & RENDERING ###

def render_graph(response: str, case: str) -> bool:
    response = str(response)
    if (response.startswith("```") and response.endswith("```")) or \
       (response.startswith('"') and (response.endswith('"') or response.endswith('"\n'))):
        try:
            dot_string = response[6:-3]
            graph = Source(dot_string, filename=f"tmp/graphs/{case}.gv", format="png")
            st.image(graph.render(), use_column_width=True, caption=f"Relaciones entre los intervinientes del caso «{case}» para los principales abonos y cargos.")
        except Exception as e:
            pass
            # st.error(f"Error al renderizar el grafo de los intervinientes: {e}")
        return True
    return False

### NARRATIVE & DOCUMENT UTILITIES ###

from datetime import datetime

def get_narrative_header(case: str, timestamp: Optional[datetime], model: str, temperature: float) -> str:
    """
    Generate a narrative header with case details, timestamp, model, and temperature.

    Args:
        case (str): Case identifier.
        timestamp (Optional[datetime]): Timestamp of the narrative generation. Defaults to current time if None.
        model (str): Name of the language model used.
        temperature (float): Temperature parameter for the model.

    Returns:
        str: A formatted narrative header.
    """
    # Use current time as fallback if timestamp is None
    if timestamp is None:
        timestamp = datetime.now()

    return (
        f"## Pre-narrativa para el caso «{case}»\n\n"
        + f"> Generada el {timestamp.strftime('%d/%m/%Y a las %H:%M')}"
        + f" usando el modelo de lenguaje «{model}» (con temperatura {temperature}).\n\n"
    )


def save_to_docx(response: str, header_image_path: str = "code/images/logo.png"):
    doc = Document()

    # Añadimos el encabezado con la imagen, si se proporciona una ruta de imagen
    if header_image_path:
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(header_image_path, width=Inches(1.0))  # Ajusta el tamaño de la imagen
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alineación a la izquierda

    lines = response.split("\n")

    # Función para identificar si una línea es parte de una tabla Markdown
    def is_table_line(line):
        return "|" in line

    # Función para procesar una tabla en formato Markdown y agregarla al documento Word
    def process_table(doc, table_lines):
        # Extraemos los encabezados de la tabla
        headers = [cell.strip() for cell in table_lines[0].split("|") if cell.strip()]
        table = doc.add_table(rows=1, cols=len(headers))

        # Aplicamos el estilo 'Table Grid' para agregar bordes
        table.style = 'Table Grid'

        # Añadimos los encabezados de la tabla
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Añadimos el resto de las filas
        for row_line in table_lines[2:]:  # Empezamos después de la fila de separadores
            row_data = [cell.strip() for cell in row_line.split("|") if cell.strip()]

            # Ajustamos el número de celdas para que coincidan con el número de encabezados
            if len(row_data) < len(headers):
                row_data += [''] * (len(headers) - len(row_data))  # Añadimos celdas vacías
            elif len(row_data) > len(headers):
                row_data = row_data[:len(headers)]  # Ignoramos celdas adicionales

            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data

    table_buffer = []  # Buffer para almacenar las líneas de la tabla

    for line in lines:
        # Ignoramos líneas vacías o con solo espacios
        if not line or line.isspace():
            continue

        # Si encontramos una línea de tabla, la almacenamos en el buffer
        if is_table_line(line):
            table_buffer.append(line)
            continue

        # Si hay un buffer de tabla y encontramos una línea que no pertenece a la tabla, procesamos la tabla
        if table_buffer and not is_table_line(line):
            process_table(doc, table_buffer)
            table_buffer = []  # Limpiamos el buffer para la siguiente tabla

        # Títulos principales: Heading nivel 0
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=0)
        # Citas: formato de Quote
        elif line.startswith("> "):
            doc.add_paragraph(line.replace("> ", ""), style='Quote')
        # Subtítulos: Heading nivel 1
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=1)
        # Párrafos normales con manejo de negrita
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Usamos un patrón para identificar palabras en negrita (dentro de **)
            pattern = r'(\*\*(.*?)\*\*|[^*]+)'
            matches = re.findall(pattern, line)
            
            for match in matches:
                bold_text = match[1]
                normal_text = match[0]
                
                if bold_text:
                    run = paragraph.add_run(bold_text)
                    run.bold = True
                else:
                    run = paragraph.add_run(normal_text)

            # Aseguramos que se añada un espacio después de cada línea procesada
            run = paragraph.add_run(" ")

    # Si el buffer de tabla aún contiene datos al final, procesamos esa tabla
    if table_buffer:
        process_table(doc, table_buffer)

    # Guardamos el documento en un stream de memoria
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

### TOKEN UTILITIES ### 

def num_tokens_from_string(string: str, encoding_name: str = "cl100k_base") -> int:
    """Returns the number of tokens in a text string."""
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens